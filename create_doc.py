"""Generate TAKUMI Trader Documentation.docx with full system documentation."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import json
from pathlib import Path
from datetime import datetime


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:color'): 'auto',
        qn('w:val'): 'clear',
    })
    tcPr.append(shading)


def add_styled_table(doc, headers, rows, col_widths=None, header_color="2E5090"):
    """Create a formatted table with colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)
                r.font.name = 'Segoe UI'
        set_cell_shading(cell, header_color)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = 'Segoe UI'
            if ri % 2 == 1:
                set_cell_shading(cell, "F0F4F8")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
        r.font.name = 'Segoe UI'
    return h


def add_para(doc, text, bold=False, italic=False, size=10, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Segoe UI'
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10)

    # ══════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TAKUMI TRADER")
    r.font.size = Pt(36)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    r.font.name = 'Segoe UI'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("匠トレーダー")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)
    r.font.name = 'Segoe UI'

    doc.add_paragraph()

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Complete System Documentation")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.name = 'Segoe UI'

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub3.add_run("All Calculations, Parameters, Settings & Trade Management")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r.font.name = 'Segoe UI'

    doc.add_paragraph()
    doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.name = 'Segoe UI'

    # Load pair settings for stats
    settings_path = Path(__file__).parent / "data" / "pair_algo_settings.json"
    pair_settings = {}
    total_bt_trades = 0
    if settings_path.exists():
        pair_settings = json.loads(settings_path.read_text(encoding="utf-8"))
        total_bt_trades = sum(
            pair_settings[p]["current"].get("trades", 0) for p in pair_settings
        )

    bt_info = doc.add_paragraph()
    bt_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = bt_info.add_run(
        f"Based on {total_bt_trades} backtested trades across 27 pairs (77-day period)"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.name = 'Segoe UI'

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "Table of Contents", 1)
    toc_items = [
        "1. System Overview",
        "2. Currency Strength Calculation Engine",
        "3. Per-Pair Algorithm Settings & Optimization",
        "4. Alert Firing Logic",
        "5. Conviction Score & Quality Filters",
        "6. Range Accumulation & Breakout Detection",
        "7. Dynamic Pip Target Calculator",
        "8. Exit Engine (5 Detectors + Trailing Stop)",
        "9. Alert Performance Tracker (MFE/MAE)",
        "10. Paper Trading System",
        "11. cTrader Auto-Trading Integration",
        "12. SL/TP Optimization Results (Per-Pair)",
        "13. Session Configuration",
        "14. ADR Range Calculation",
        "15. All Configurable Settings",
    ]
    for item in toc_items:
        add_para(doc, item, size=11, space_after=4)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. SYSTEM OVERVIEW
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. System Overview", 1)
    add_para(doc, (
        "TAKUMI Trader (匠 — master craftsman) is a real-time currency strength scanner and "
        "algorithmic execution engine that monitors all 28 major forex pairs across 4 timeframes "
        "(M1, M5, M15, H1). It calculates relative currency strength for 8 major currencies "
        "(USD, EUR, GBP, JPY, CAD, AUD, NZD, CHF), identifies the strongest divergences, and "
        "fires alerts when conditions align."
    ))
    add_para(doc, (
        "The system includes automatic trade tracking with MFE/MAE analysis, a 5-detector exit engine, "
        "conviction scoring with 4 quality filters, range accumulation detection, per-pair optimized "
        "SL/TP parameters, a local paper trading engine, and optional auto-trading via cTrader Open API."
    ))

    add_heading_styled(doc, "Architecture", 2)
    arch_rows = [
        ["Strength Engine", "core/strength.py", "EMA displacement + weighted ROC + tick velocity per currency"],
        ["Alert Manager", "core/alerts.py", "Cooldown management, notification dispatch"],
        ["Filter Engine", "core/filter_engine.py", "4 quality filters, conviction scoring (0-100)"],
        ["Range Engine", "core/range_engine.py", "Range accumulation, BB squeeze, breakout detection"],
        ["Target Calculator", "core/target_calculator.py", "Dynamic pip targets based on ATR, session, ADR, conviction"],
        ["Exit Engine", "core/exit_engine.py", "5 exit detectors + trailing strength stop"],
        ["Trade Tracker", "core/trade_tracker.py", "P/L tracking, peak/worst, partial targets, paper trade fields"],
        ["Alert Performance", "core/alert_performance.py", "MFE/MAE tracking, 2-phase observation, post-exit analysis"],
        ["Paper Trader", "core/paper_trader.py", "Local paper trading with SL/TP management and journaling"],
        ["Pair Algo Settings", "core/pair_algo_settings.py", "Per-pair optimized parameters with history"],
        ["HTF Regime", "core/htf_regime.py", "H4/D1 trend regime tracking, velocity measurement"],
        ["Session Manager", "core/session_manager.py", "DST-aware session detection using ZoneInfo"],
        ["MT5 Worker", "core/mt5_worker.py", "MetaTrader 5 data feed thread (M1/M5/M15/H1/H4/D1)"],
        ["cTrader Bridge", "core/ctrader_worker.py", "Auto-trading via cTrader Open API"],
        ["cTrader Positions", "core/ctrader_position_manager.py", "Position state tracking & duplicate prevention"],
    ]
    add_styled_table(doc, ["Component", "File", "Purpose"], arch_rows, [4, 5, 8])

    add_heading_styled(doc, "Technology Stack", 2)
    tech_rows = [
        ["Framework", "PyQt6 native Windows 11 desktop application"],
        ["Data Feed", "MetaTrader 5 (live M1/M5/M15/H1/H4/D1 candles)"],
        ["Historical Data", "Dukascopy (M1 data for backtesting)"],
        ["Execution", "cTrader Open API v2 via Protobuf/TCP (Twisted reactor)"],
        ["Performance", "Numba JIT-compiled indicators for low-latency calculations"],
        ["Notifications", "Windows 11 Toast notifications + WAV/MP3 sound alerts"],
    ]
    add_styled_table(doc, ["Component", "Description"], tech_rows, [4, 13])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. CURRENCY STRENGTH CALCULATION ENGINE
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Currency Strength Calculation Engine", 1)

    add_heading_styled(doc, "Core Parameters", 2)
    core_rows = [
        ["Currencies", "USD, EUR, GBP, JPY, CAD, AUD, NZD, CHF", "8 major currencies"],
        ["Total Pairs", "28", "All combinations of the 8 currencies"],
        ["Timeframes", "M1, M5, M15, H1", "4 concurrent timeframes"],
        ["Warmup Bars", "200", "Initial bars needed for calculations"],
        ["Z-Score Window", "120 bars", "Rolling buffer for Z-score normalization"],
        ["EMA Period", "6 or 8 bars", "Per-pair optimized (see Section 3)"],
        ["ATR Period", "14 bars", "Average True Range for normalization"],
        ["ROC Decay", "0.2 or 0.3", "Per-pair optimized exponential decay"],
        ["Tanh Sensitivity", "1.0", "Controls output compression"],
        ["Output Range", "±10", "Normalized strength range after tanh compression"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], core_rows, [4, 5, 8])

    add_heading_styled(doc, "ROC Lookback Periods", 2)
    roc_rows = [
        ["M1", "10", "150"],
        ["M5", "8", "50"],
        ["M15", "6", "50"],
        ["H1", "5", "50"],
    ]
    add_styled_table(doc, ["Timeframe", "Lookback Bars", "Live Fetch Bars"], roc_rows, [4, 4, 4])

    add_heading_styled(doc, "Composite Score Weights", 2)
    weight_rows = [
        ["M1", "0.35", "0.35", "0.30"],
        ["M5", "0.50", "0.50", "0.00"],
        ["M15", "0.50", "0.50", "0.00"],
        ["H1", "0.50", "0.50", "0.00"],
    ]
    add_styled_table(doc, ["Timeframe", "Displacement", "Weighted ROC", "Tick Velocity"], weight_rows)

    add_heading_styled(doc, "Calculation Formulas", 2)
    formulas = [
        ("EMA-8 Displacement", "displacement = (close - EMA8) / ATR14"),
        ("Weighted Micro-ROC", "For each i in range(lookback): weight = exp(-decay * i); roc = (close - close[i ago]) / ATR14; result = sum(weight * roc) / sum(weights)"),
        ("Tick Velocity (M1 only)", "tick_velocity = (current_close - candle_open) / ATR14"),
        ("ATR14", "TR = max(H-L, |H-prevC|, |L-prevC|); ATR = EMA(TR, 14)"),
        ("Per-Currency Raw Score", "raw[ccy] = average of signed pair scores for all 7 pairs containing that currency"),
        ("Z-Score Normalization", "z = (raw - mean) / std_dev; normalized = 10.0 * tanh(z * sensitivity)"),
        ("Pair Display Score", "pair_score = (base_strength - quote_strength) / 2.0"),
        ("Momentum Acceleration", "accel = (score_now - score_prev) - (score_prev - score_2ago)"),
    ]
    for name, formula in formulas:
        p = doc.add_paragraph()
        r = p.add_run(f"{name}: ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = 'Segoe UI'
        r2 = p.add_run(formula)
        r2.font.size = Pt(9)
        r2.font.name = 'Consolas'

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. PER-PAIR ALGORITHM SETTINGS & OPTIMIZATION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Per-Pair Algorithm Settings & Optimization", 1)
    add_para(doc, (
        "Each of the 27 traded pairs has individually optimized calculation parameters determined "
        "via a 3-phase backtest optimization process over a 77-day period (January 5 - March 23, 2026). "
        "Settings are stored in data/pair_algo_settings.json with version history."
    ))

    add_heading_styled(doc, "Optimization Methodology", 2)
    opt_phases = [
        ["Phase 1", "Grid Search", "Sweeps ema_period, roc_decay, sensitivity, threshold_m1, min_divergence_spread, require_acceleration"],
        ["Phase 2", "Threshold Sweep", "Post-hoc threshold tuning on filtered outcomes"],
        ["Phase 3", "SL/TP Optimization", "ATR-multiplier based SL/TP simulation with hybrid/strict modes"],
    ]
    add_styled_table(doc, ["Phase", "Method", "Description"], opt_phases, [3, 4, 10])

    add_heading_styled(doc, "Per-Pair Optimized Parameters", 2)
    if pair_settings:
        pair_rows = []
        for p in sorted(pair_settings.keys()):
            c = pair_settings[p]["current"]
            pair_rows.append([
                p,
                str(c.get("ema_period", "")),
                str(c.get("roc_decay", "")),
                f'{c.get("sl_pips", 0):.1f}',
                f'{c.get("tp_pips", 0):.1f}',
                f'{c.get("sl_atr", 0)}',
                f'{c.get("tp_atr", 0)}',
                f'{c.get("wr", 0):.1f}%',
                str(c.get("trades", "")),
            ])
        add_styled_table(
            doc,
            ["Pair", "EMA", "ROC Decay", "SL (pips)", "TP (pips)", "SL ATR", "TP ATR", "Win Rate", "Trades"],
            pair_rows,
            [2.5, 1.3, 1.8, 1.8, 1.8, 1.5, 1.5, 1.8, 1.5],
        )

    add_heading_styled(doc, "Common Optimized Parameters", 2)
    common_rows = [
        ["threshold_m1", "5.5", "M1 strength threshold for alert firing"],
        ["min_divergence_spread", "12.0", "Minimum spread between base and quote scores"],
        ["require_acceleration", "False", "Acceleration not required for alert"],
        ["sensitivity", "1.0", "Tanh compression sensitivity"],
        ["sl_atr (most pairs)", "0.3", "Stop loss = 0.3 × ATR14"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], common_rows, [5, 3, 9])

    add_para(doc, (
        "Settings include version history (last 50 entries) and can be restored via the "
        "Pair Algorithm Settings dialog accessible from Settings > Pair Algorithm Settings."
    ), italic=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. ALERT FIRING LOGIC
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Alert Firing Logic", 1)
    alert_rows = [
        ["Default Cooldown", "60 seconds", "Min time between alerts for same pair"],
        ["Cooldown Range", "1-30 minutes", "User configurable in settings"],
        ["Trigger Condition", "New pair enters alert state", "Only fires for newly qualifying pairs"],
        ["Notification", "Windows 11 Toast + Sound", "Optional WAV/MP3 playback"],
        ["Conviction Filter", "FULL tier only", "Only fires with sound for conviction ≥ 70"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], alert_rows, [4, 5, 8])

    add_para(doc, (
        "Alert format: \"{PAIR} - {DIRECTION} (M1: +X.X | M5: +X.X | M15: +X.X)\". "
        "Alerts fire when a pair newly enters the alert state based on strong currency strength "
        "divergence. DIMMED alerts (conviction 45-69) are shown visually but without sound or "
        "auto-trading. SUPPRESSED alerts (conviction < 45) are hidden entirely."
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. CONVICTION SCORE & QUALITY FILTERS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Conviction Score & Quality Filters", 1)
    add_para(doc, (
        "The conviction score (0-100) is computed from 4 independent quality filters. Each filter "
        "contributes points to a raw total of 85, which is normalized to 0-100. Disabled filters "
        "contribute full points (neutral). Each filter can be toggled on/off via the UI toolbar."
    ))

    add_heading_styled(doc, "Filter 1: HTF Trend Regime (30 pts max)", 2)
    htf_rows = [
        ["H4 Strong Currency Aligned", "8", "Strong ccy is BULLISH (BUY) or BEARISH (SELL) on H4"],
        ["H4 Weak Currency Aligned", "8", "Weak ccy is BEARISH (BUY) or BULLISH (SELL) on H4"],
        ["H4 Fully Aligned Bonus", "4", "Both H4 conditions met"],
        ["D1 Strong Currency Aligned", "4", "Strong ccy aligned on D1"],
        ["D1 Weak Currency Aligned", "4", "Weak ccy aligned on D1"],
        ["D1 Fully Aligned Bonus", "2", "Both D1 conditions met"],
        ["Regime Threshold", "3.0", "H4/D1 strength required for BULL/BEAR classification"],
    ]
    add_styled_table(doc, ["Component", "Points", "Condition"], htf_rows, [5, 2, 10])

    add_heading_styled(doc, "Filter 2: Strength Velocity (20 pts max)", 2)
    vel_rows = [
        ["Velocity Threshold", "0.6 pts/min", "Speed threshold for 'fast' classification"],
        ["Velocity Max Scale", "1.2", "Velocity at which full points are awarded"],
        ["Strong Currency Score", "0-10 pts", "Based on abs(velocity) / max_scale * 10"],
        ["Weak Currency Score", "0-10 pts", "Same formula for weak currency"],
        ["Pass Threshold", "8 pts", "Minimum combined score to 'pass'"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], vel_rows, [5, 3, 9])

    add_heading_styled(doc, "Filter 3: Isolation Score (20 pts max)", 2)
    iso_rows = [
        ["Strong #1 + Big Gap (≥ 2.0)", "8", "Strongest currency with large gap to #2"],
        ["Strong #1 + Small Gap (≥ 1.0)", "5", "Strongest but gap is moderate"],
        ["Strong #1 + Bunched", "2", "Strongest but too close to #2"],
        ["Weak #8 + Big Gap (≥ 2.0)", "8", "Weakest currency with large gap to #7"],
        ["Weak #8 + Small Gap (≥ 1.0)", "5", "Weakest but gap is moderate"],
        ["Weak #8 + Bunched", "2", "Weakest but too close to #7"],
        ["#1 vs #8 Bonus", "4", "Trading strongest against weakest"],
    ]
    add_styled_table(doc, ["Component", "Points", "Condition"], iso_rows, [5, 2, 10])

    add_heading_styled(doc, "Filter 4: ADR Position (15 pts max)", 2)
    adr_rows = [
        ["< 40%", "15", "ADR Fresh"],
        ["40-60%", "12", "ADR Healthy"],
        ["60-75%", "8", "ADR Moderate"],
        ["75-85%", "4", "ADR Limited"],
        ["≥ 85%", "0", "ADR Exhausted"],
    ]
    add_styled_table(doc, ["ADR Consumed", "Points", "Label"], adr_rows, [4, 3, 5])

    add_heading_styled(doc, "Conviction Tiers", 2)
    tier_rows = [
        ["FULL", "≥ 70", "Full alert with sound + auto-trading + paper trade eligible"],
        ["DIMMED", "45-69", "Alert fires but dimmed visual, no auto-trade"],
        ["SUPPRESSED", "< 45", "Alert suppressed entirely"],
    ]
    add_styled_table(doc, ["Tier", "Threshold", "Behavior"], tier_rows, [3, 3, 11])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. RANGE ACCUMULATION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Range Accumulation & Breakout Detection", 1)
    range_rows = [
        ["Window Size", "10 M1 candles", "Lookback for range detection"],
        ["Range Threshold", "5.0%", "Max range as % of ADR for 'in range'"],
        ["Min Quality", "60", "Quality score needed for LOADED tier"],
        ["Strength Delta Min", "3.0", "Min delta for direction prediction"],
        ["Breakout Cooldown", "300 seconds", "5 min cooldown after breakout"],
        ["Max ADR Consumed", "85%", "Suppress if daily range spent"],
        ["BB Period", "20", "Bollinger Band period"],
        ["BB Width", "4 std dev", "Bollinger Band width (2 std × 2 sides)"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], range_rows, [4, 4, 9])

    add_heading_styled(doc, "Range Tiers", 2)
    rtier_rows = [
        ["RANGE", "Blue", "Range detected (visual only)"],
        ["LOADED", "Purple", "Range + quality ≥ 60 + strength building"],
        ["BREAKOUT", "Cyan", "Price broke range in predicted direction"],
        ["PRIME", "Gold", "Strength FULL ALERT + LOADED aligned (highest conviction)"],
        ["STALE", "Gray", "Range persisted > 2× window without breakout"],
    ]
    add_styled_table(doc, ["Tier", "Color", "Condition"], rtier_rows, [3, 3, 11])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. DYNAMIC PIP TARGET
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. Dynamic Pip Target Calculator", 1)
    target_rows = [
        ["JPY Pairs", "12.0 pips", "5.0 pips", "35.0 pips"],
        ["Non-JPY Pairs", "8.0 pips", "3.0 pips", "25.0 pips"],
    ]
    add_styled_table(doc, ["Pair Type", "Base Target", "Min", "Max"], target_rows)

    add_heading_styled(doc, "Session Multipliers", 2)
    sess_mult_rows = [
        ["Tokyo", "0.7", "Lower volatility"],
        ["Frankfurt", "0.9", "Moderate volatility"],
        ["London", "1.0", "High volatility (baseline)"],
        ["Overlap (London+NY)", "1.2", "Highest volatility"],
        ["NY PM", "0.8", "Declining volatility"],
        ["Off Hours", "0.5", "Minimal volatility"],
    ]
    add_styled_table(doc, ["Session", "Multiplier", "Description"], sess_mult_rows, [4, 3, 10])

    add_heading_styled(doc, "Target Formula", 2)
    add_para(doc, "target = base × atr_factor × adr_factor × session_factor × strength_factor × conviction_mult", size=9)
    add_para(doc, "Conviction multiplier: conv_mult = 0.7 + (conviction/100) × 0.6  [range: 0.7 to 1.3]", size=9)
    add_para(doc, "Partial target = 50% of full target", size=9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. EXIT ENGINE
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. Exit Engine (5 Detectors + Trailing Stop)", 1)
    add_para(doc, (
        "Five independent exit signal detectors each cast a boolean vote. Votes are aggregated "
        "into urgency levels. A trailing strength stop can override the vote count upward."
    ))

    add_heading_styled(doc, "Detector 1: Strength Reversal", 2)
    d1_rows = [
        ["Base Drop Threshold", "3.0 points", "Base currency dropped this much from entry"],
        ["Quote Rise Threshold", "3.0 points", "Quote currency rose this much from entry"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], d1_rows, [5, 3, 9])

    add_heading_styled(doc, "Detector 2: Momentum Stall", 2)
    d2_rows = [
        ["Flat Threshold", "0.3", "M1 score change < this = flat"],
        ["Reverse Threshold", "-1.0", "M1 pair score moved against trade"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], d2_rows, [5, 3, 9])

    add_heading_styled(doc, "Detector 3: Range Exhaustion", 2)
    add_para(doc, "ADR Exhaustion Threshold: 80% of daily range consumed")

    add_heading_styled(doc, "Detector 4: Time Decay", 2)
    d4_rows = [
        ["Base Scalp Limit", "15.0 minutes", "Default time limit"],
        ["High Conviction (≥80)", "19.5 minutes", "15 × 1.3 multiplier"],
        ["Medium Conviction (60-79)", "15.0 minutes", "15 × 1.0 multiplier"],
        ["Low Conviction (<60)", "10.5 minutes", "15 × 0.7 multiplier"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], d4_rows, [5, 3, 9])

    add_heading_styled(doc, "Detector 5: Adverse Flow", 2)
    add_para(doc, "Flow Threshold: -0.15 (tick flow bias range is -1 to +1)")

    add_heading_styled(doc, "Trailing Strength Stop", 2)
    ts_rows = [
        ["Drop Percentage", "40%", "Strength delta dropped 40% from peak"],
        ["Minimum Peak", "3.0", "Peak delta must be ≥ 3.0"],
        ["Override", "Forces vote_count to 3 (CLOSE)", "If trailing stop triggered and votes < 3"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], ts_rows, [4, 5, 8])

    add_heading_styled(doc, "Vote Escalation", 2)
    vote_rows = [
        ["0-1", "(none)", "(none)"],
        ["2", "WATCH", "TIGHTEN stop mentally"],
        ["3", "CLOSE", "EXIT — consider closing"],
        ["4-5", "URGENT", "EXIT — close immediately"],
    ]
    add_styled_table(doc, ["Votes", "Urgency", "Action"], vote_rows, [3, 4, 10])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 9. ALERT PERFORMANCE TRACKER
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Alert Performance Tracker (MFE/MAE)", 1)

    add_heading_styled(doc, "Tracking Parameters", 2)
    perf_rows = [
        ["Exit Spread Threshold", "4.0", "Strength spread must drop below this to trigger exit"],
        ["Post-Exit Observation", "4.0 hours", "Continue tracking after exit signal"],
        ["Max Tracking Hours", "8.0 hours", "Safety timeout if no exit signal fires"],
        ["Gap Backfill Min", "120 seconds", "Minimum gap to trigger backfill"],
        ["Gap Backfill Max", "1500 bars", "Max M1 bars to fetch for backfill"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], perf_rows, [5, 3, 9])

    add_heading_styled(doc, "Two-Phase Tracking System", 2)
    add_para(doc, "Phase 1 (Entry → Exit Signal): Tracks MFE/MAE in the alert direction. Records time to reach MFE and MAE. Monitors currency strength spread for exit signal.", bold=False)
    add_para(doc, "Phase 2 (Post-Exit +4 hours): After exit signal fires, continues tracking for 4 more hours. Records post-exit MFE/MAE relative to exit point. Determines final P/L at end of observation.")

    add_heading_styled(doc, "HTF Exit Signal", 2)
    add_para(doc, (
        "Exit signals use HTF composite scores (M5 + M15 + H1 average, excluding M1) to filter out "
        "M1 noise and pullbacks. This prevents premature exits during temporary M1 retracements, "
        "allowing trades to run through minor pullbacks."
    ))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 10. PAPER TRADING SYSTEM
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. Paper Trading System", 1)
    add_para(doc, (
        "The Paper Trading system simulates full trade execution locally, providing an execution "
        "benchmark for comparison with live broker fills (cTrader, Interactive Brokers). All FULL "
        "conviction alerts automatically open paper trades with per-pair optimized SL/TP levels."
    ))

    add_heading_styled(doc, "How It Works", 2)
    paper_steps = [
        ["1. Alert Fires", "FULL conviction alert → Paper trade auto-opens at signal close price"],
        ["2. SL/TP Set", "Looks up per-pair optimized sl_pips/tp_pips from pair_algo_settings.json"],
        ["3. Price Levels", "Computes exact SL/TP price levels based on entry price and direction"],
        ["4. M1 Monitoring", "Every M1 cycle checks candle High/Low against SL/TP levels"],
        ["5. Exit Detection", "Closes on: SL hit (Low ≤ SL for BUY), TP hit (High ≥ TP for BUY), or URGENT signal"],
        ["6. Journal", "Records complete trade record to data/paper_trades.json"],
        ["7. Performance", "Stats available in Performance dialog → Paper Trades tab"],
    ]
    add_styled_table(doc, ["Step", "Description"], paper_steps, [3, 14])

    add_heading_styled(doc, "SL/TP Hit Detection (M1 High/Low)", 2)
    add_para(doc, (
        "SL/TP detection uses M1 candle High and Low prices (not just the Close) for realistic "
        "simulation. For BUY trades: TP hit if candle High ≥ TP price, SL hit if candle Low ≤ SL price. "
        "For SELL trades: TP hit if candle Low ≤ TP price, SL hit if candle High ≥ SL price. "
        "TP is checked first (optimistic fill assumption)."
    ))

    add_heading_styled(doc, "Paper Trade Record Fields", 2)
    paper_fields = [
        ["pair, direction", "Trade identity"],
        ["entry_price, entry_time", "Entry point and timestamp"],
        ["close_price, close_time", "Exit point and timestamp"],
        ["close_reason", "sl_hit / tp_hit / signal_exit"],
        ["sl_pips, tp_pips", "SL/TP distances from pair_algo_settings"],
        ["sl_price, tp_price", "Computed absolute price levels"],
        ["pnl_pips", "Final P/L in pips"],
        ["peak_pnl_pips, worst_pnl_pips", "Best and worst P/L during trade"],
        ["duration_minutes", "Total trade duration"],
        ["entry_conviction", "Conviction score at entry"],
        ["session", "Trading session at entry"],
    ]
    add_styled_table(doc, ["Field", "Description"], paper_fields, [5, 12])

    add_heading_styled(doc, "Dead Zone", 2)
    add_para(doc, (
        "No new paper trades open between 5:01-6:59 JST (early morning dead zone). "
        "Active paper trades continue to be monitored for SL/TP hits during this period."
    ))

    add_heading_styled(doc, "Execution Benchmark Purpose", 2)
    add_para(doc, (
        "Paper trades serve as the ideal execution benchmark — instant fills at exact SL/TP levels "
        "with zero latency. When broker trades come in (via cTrader or future Interactive Brokers "
        "integration), they can be compared against paper trades to measure slippage, fill latency, "
        "SL/TP drift, and missed trades."
    ))

    add_heading_styled(doc, "Data Storage", 2)
    storage_rows = [
        ["data/paper_trades.json", "Completed paper trade journal (append-mode)"],
        ["data/tracked_trades.json", "Active paper trades (survive restart via is_paper flag)"],
    ]
    add_styled_table(doc, ["File", "Purpose"], storage_rows, [6, 11])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 11. CTRADER AUTO-TRADING
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. cTrader Auto-Trading Integration", 1)
    ct_rows = [
        ["Protocol", "Protobuf over TCP", "Port 5035"],
        ["Auth", "OAuth2", "Client ID + Secret + Access Token"],
        ["Heartbeat", "10-20 seconds", "Keep-alive ping"],
        ["Volume Conversion", "1 lot = 10,000,000 cents", "volume_cents = lots × 100,000 × 100"],
        ["Order Type", "MARKET (type 1)", "Immediate execution"],
        ["Reconnect Backoff", "5s to 60s max", "Exponential with cap"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], ct_rows, [4, 5, 8])

    add_heading_styled(doc, "Auto-Trading Rules", 2)
    at_rows = [
        ["Auto-Open", "Opens market order on FULL conviction alerts only"],
        ["Auto-Close", "Closes position on URGENT exit signal"],
        ["Duplicate Prevention", "No double positions on same pair+direction"],
        ["Max Positions", "1-10 (configurable, default 3)"],
        ["Lot Size", "0.01-1.0 (configurable, default 0.01)"],
        ["Dead Zone", "No new trades 5:01-6:59 JST"],
    ]
    add_styled_table(doc, ["Rule", "Description"], at_rows, [4, 13])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 12. SL/TP OPTIMIZATION RESULTS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "12. SL/TP Optimization Results (Per-Pair)", 1)
    add_para(doc, (
        f"Based on {total_bt_trades} backtested trades across 27 pairs over a 77-day period "
        f"(January 5 - March 23, 2026). Each pair has individually optimized ATR-based SL/TP "
        f"multipliers determined by the 3-phase optimizer."
    ))

    add_heading_styled(doc, "ATR Multiplier Approach", 2)
    add_para(doc, (
        "Stop Loss and Take Profit levels are set as multiples of the H1 ATR(14) at the time of "
        "entry. This automatically adapts to each pair's current volatility. Most pairs converged "
        "on SL ATR = 0.3 (tight stop), while TP ATR ranges from 0.5 to 3.0 depending on the "
        "pair's typical move characteristics."
    ))

    if pair_settings:
        add_heading_styled(doc, "Full Per-Pair Results", 2)
        full_rows = []
        for p in sorted(pair_settings.keys()):
            c = pair_settings[p]["current"]
            full_rows.append([
                p,
                f'{c.get("sl_pips", 0):.1f}',
                f'{c.get("tp_pips", 0):.1f}',
                f'{c.get("wr", 0):.1f}%',
                f'{c.get("exp_r", 0):.2f}',
                f'{c.get("total_r", 0):.1f}',
                f'{c.get("avg_mfe", 0):.1f}',
                f'{c.get("avg_mae", 0):.1f}',
                str(c.get("trades", "")),
            ])
        add_styled_table(
            doc,
            ["Pair", "SL (pips)", "TP (pips)", "WR", "Exp R", "Total R", "Avg MFE", "Avg MAE", "Trades"],
            full_rows,
            [2.5, 1.8, 1.8, 1.5, 1.5, 1.8, 1.8, 1.8, 1.5],
        )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 13. SESSION CONFIGURATION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "13. Session Configuration", 1)
    add_para(doc, (
        "Sessions are defined using local city times and converted to UTC at runtime using Python's "
        "zoneinfo module. This means DST transitions are handled automatically — no hardcoded UTC offsets."
    ))
    sess_rows = [
        ["Tokyo", "Asia/Tokyo", "09:00-15:00", "JPY pairs most active, lower volatility"],
        ["Frankfurt", "Europe/Berlin", "08:00-12:00", "EUR pre-London, moderate"],
        ["London", "Europe/London", "08:00-12:00", "GBP/EUR pairs, high volatility"],
        ["Overlap", "America/New_York", "08:00-12:00", "London+NY, highest volatility"],
        ["NY PM", "America/New_York", "12:00-17:00", "USD pairs, declining volatility"],
        ["Off Hours", "—", "—", "Between sessions, minimal activity"],
    ]
    add_styled_table(doc, ["Session", "Timezone", "Local Hours", "Characteristics"], sess_rows, [3, 4, 3, 7])

    add_para(doc, (
        "Dead Zone: No new trades between 5:01-6:59 JST. Alert tracking and performance "
        "monitoring continue during this period but no paper trades or cTrader orders are opened."
    ), italic=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 14. ADR RANGE CALCULATION
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "14. ADR Range Calculation", 1)
    add_para(doc, (
        "The Average Daily Range (ADR) measures how much a pair typically moves in a day. "
        "It is used throughout the system for range exhaustion detection, target calculation, "
        "and conviction scoring."
    ))

    adr_params = [
        ["ADR Period", "14 days", "Rolling 14-day average of daily High-Low range"],
        ["Day Boundary", "07:00 JST", "Trading day starts at 7:00 AM JST (≈22:00 UTC)"],
        ["Range Scope", "24 hours", "Full 24-hour daily range (all sessions)"],
        ["Display", "Range column (%)", "Current day's range as percentage of ADR"],
    ]
    add_styled_table(doc, ["Parameter", "Value", "Description"], adr_params, [4, 4, 9])

    add_para(doc, (
        "The 07:00 JST day boundary was chosen because it falls in the early morning dead zone "
        "between the NY close and the Tokyo open, avoiding splitting any active session's range "
        "across two days. This aligns closely with the standard forex daily candle rollover time "
        "(~22:00 UTC)."
    ))

    add_heading_styled(doc, "ADR Usage in the System", 2)
    adr_usage = [
        ["Filter Engine (ADR Position)", "Scores conviction based on remaining daily range (0-15 pts)"],
        ["Exit Engine (Range Exhaustion)", "Triggers exit vote when >80% of ADR consumed"],
        ["Target Calculator (ADR Remaining)", "Reduces pip target when range is mostly consumed"],
        ["Range Engine (Range Threshold)", "Detects accumulation when range < 5% of ADR"],
        ["Range Column (UI)", "Shows current consumed % with color coding"],
    ]
    add_styled_table(doc, ["Usage", "Description"], adr_usage, [5, 12])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 15. ALL CONFIGURABLE SETTINGS
    # ══════════════════════════════════════════════════════════════
    add_heading_styled(doc, "15. All Configurable Settings", 1)

    add_heading_styled(doc, "Alert Settings (UI)", 2)
    alert_settings = [
        ["Sound Enabled", "True", "On/Off", "Settings dialog"],
        ["Sound File", "(none)", "WAV/MP3 path", "Settings dialog"],
        ["Alert Cooldown", "60 seconds", "1-30 minutes", "Settings dialog"],
    ]
    add_styled_table(doc, ["Setting", "Default", "Range", "Location"], alert_settings)

    add_heading_styled(doc, "UI Settings", 2)
    ui_settings = [
        ["Font Size", "10 pt", "7-18 pt"],
        ["Compact Mode", "Off", "On/Off"],
        ["Pin on Top", "Off", "On/Off"],
    ]
    add_styled_table(doc, ["Setting", "Default", "Range"], ui_settings)

    add_heading_styled(doc, "Filter Toggles (Toolbar)", 2)
    filter_toggles = [
        ["HTF", "On", "HTF Trend Regime filter (30 pts)"],
        ["VEL", "On", "Strength Velocity filter (20 pts)"],
        ["ISOL", "On", "Isolation Score filter (20 pts)"],
        ["ADR", "On", "ADR Position filter (15 pts)"],
        ["Conv ≥ 70", "On", "Minimum conviction threshold for display"],
    ]
    add_styled_table(doc, ["Button", "Default", "Description"], filter_toggles, [3, 3, 11])

    add_heading_styled(doc, "cTrader Settings (UI)", 2)
    ct_settings = [
        ["Enabled", "Off", "On/Off"],
        ["Client ID", "(empty)", "From cTrader Open API portal"],
        ["Client Secret", "(empty)", "Masked input"],
        ["Access Token", "(empty)", "Masked input"],
        ["Account ID", "(empty)", "Demo account number"],
        ["Lot Size", "0.01", "0.01-1.0 (step 0.01)"],
        ["Auto-Open", "On", "On/Off"],
        ["Auto-Close", "On", "On/Off"],
        ["Max Positions", "3", "1-10"],
    ]
    add_styled_table(doc, ["Setting", "Default", "Range"], ct_settings)

    add_heading_styled(doc, "Internal Constants (Code)", 2)
    const_rows = [
        ["EXIT_SPREAD_THRESHOLD", "4.0", "alert_performance.py"],
        ["POST_EXIT_HOURS", "4.0", "alert_performance.py"],
        ["MAX_TRACKING_HOURS", "8.0", "alert_performance.py"],
        ["STRENGTH_REVERSAL_DROP", "3.0", "exit_engine.py"],
        ["MOMENTUM_REVERSE_THRESHOLD", "-1.0", "exit_engine.py"],
        ["ADR_EXHAUSTION_PCT", "80.0", "exit_engine.py"],
        ["MAX_SCALP_MINUTES", "15.0", "exit_engine.py"],
        ["FLOW_AGAINST_THRESHOLD", "-0.15", "exit_engine.py"],
        ["TRAILING_STRENGTH_DROP_PCT", "40.0", "exit_engine.py"],
        ["CONVICTION_FULL_THRESHOLD", "70", "filter_engine.py"],
        ["CONVICTION_DIMMED_THRESHOLD", "45", "filter_engine.py"],
        ["ZSCORE_WINDOW", "120", "strength.py"],
        ["WARMUP_BARS", "200", "strength.py"],
        ["RANGE_THRESHOLD_PCT", "5.0", "range_engine.py"],
        ["RANGE_WINDOW_SIZE", "10", "range_engine.py"],
        ["MIN_QUALITY", "60", "range_engine.py"],
        ["BREAKOUT_COOLDOWN", "300s", "range_engine.py"],
        ["DEFAULT_SL_PIPS", "10.0", "paper_trader.py (fallback)"],
        ["DEFAULT_TP_PIPS", "20.0", "paper_trader.py (fallback)"],
    ]
    add_styled_table(doc, ["Constant", "Value", "File"], const_rows, [6, 3, 6])

    add_heading_styled(doc, "Data Storage Files", 2)
    data_rows = [
        ["data/pair_algo_settings.json", "Per-pair optimized parameters + history"],
        ["data/alert_outcomes.json", "Live alert tracking results (MFE/MAE)"],
        ["data/backtest_outcomes.json", "Raw backtest trade results"],
        ["data/paper_trades.json", "Completed paper trade journal"],
        ["data/active_perf_alerts.json", "Currently tracking alerts (crash recovery)"],
        ["data/tracked_trades.json", "Active tracked/paper trades (crash recovery)"],
        ["data/ctrader_positions.json", "Open cTrader positions"],
        ["data/alert_history.json", "Historical alert records"],
        ["data/optimal_sltp.json", "Session-specific fixed-pip SL/TP profiles"],
        ["data/optimal_sltp_atr.json", "ATR-based SL/TP profiles"],
    ]
    add_styled_table(doc, ["File", "Purpose"], data_rows, [6, 11])

    # ── Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("— End of Documentation —")
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    r.font.name = 'Segoe UI'

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = note.add_run("This application is intended for private, personal use and is not for commercial distribution.")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r.font.name = 'Segoe UI'

    # Save
    output_path = Path(__file__).parent / "TAKUMI Trader Documentation.docx"
    doc.save(str(output_path))
    print(f"Documentation saved to: {output_path}")
    print(f"Sections: 15")
    print(f"Pairs documented: {len(pair_settings)}")
    print(f"Backtest trades: {total_bt_trades}")


if __name__ == "__main__":
    main()
